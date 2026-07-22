# -*- coding: utf-8 -*-
"""Rellena plantillas Word/Excel con datos de la empresa activa.

Las plantillas del pack traen datos de ejemplo (FOREST GREEN, NIT demo, etc.).
Se reemplazan a nivel de parrafo y tambien en el XML interno del .docx/.xlsx
(para atrapar texto partido en varios runs / cuadros de texto).
"""
from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path
from typing import Any, Iterable
from xml.sax.saxutils import escape as xml_escape

from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage

from generador.familias_documento import contexto_actividad

_EMPRESA_EJEMPLO_REGEX = [
    re.compile(r"FOREST\s+GREEN\s+SERVICIOS\s+AMBIENTALES\s+S\.?\s*A\.?\s*S\.?", re.I),
    re.compile(r"FOREST\s+GREEN\s+SERVICIOS\s+AMBIENTALES", re.I),
    re.compile(r"\bFOREST\s+GREEN\b", re.I),
    re.compile(r"Laboratorios\s+LT\s+S\.?\s*A\.?\s*S\.?", re.I),
    re.compile(r"\bXXXXXX(?:\s+XXXXXX)+\b", re.I),
    re.compile(r"\[?\s*NOMBRE\s+DE\s+LA\s+EMPRESA\s*\]?", re.I),
    re.compile(r"\[?\s*RAZ[O\u00d3]N\s+SOCIAL\s*\]?", re.I),
]

_NIT_EJEMPLO_REGEX = [
    re.compile(r"\b901[\.\s]?989[\.\s]?693\s*-?\s*7\b"),
    re.compile(r"\b901989693-?7\b"),
]


def _ctx_val(ctx: dict[str, str], key: str, fallback: str = "________________") -> str:
    v = (ctx.get(key) or "").strip()
    return v if v else fallback


def _contexto(empresa: dict, doc_meta: dict, version: int, fecha: str) -> dict[str, str]:
    clases = ", ".join(empresa.get("clases_riesgo_list") or []) or (
        empresa.get("clases_riesgo") or ""
    )
    return {
        "razon_social": (empresa.get("razon_social") or "").strip(),
        "empresa": (empresa.get("razon_social") or "").strip(),
        "nit": (empresa.get("nit") or "").strip(),
        "arl": (empresa.get("arl") or "").strip(),
        "ciiu": f"{empresa.get('ciiu_codigo', '')} {empresa.get('ciiu_descripcion', '')}".strip(),
        "ciiu_codigo": (empresa.get("ciiu_codigo") or "").strip(),
        "ciiu_descripcion": (empresa.get("ciiu_descripcion") or "").strip(),
        "actividad": contexto_actividad(empresa),
        "clases_riesgo": clases,
        "codigo": doc_meta.get("codigo", "") or "",
        "version": str(version),
        "fecha": fecha,
        "titulo": doc_meta.get("nombre", "") or "",
        "nombre_documento": doc_meta.get("nombre", "") or "",
        "rep_legal_nombre": (empresa.get("rep_legal_nombre") or "").strip(),
        "rep_legal_cargo": (empresa.get("rep_legal_cargo") or "").strip(),
        "resp_sst_nombre": (empresa.get("resp_sst_nombre") or "").strip(),
        "resp_sst_cargo": (empresa.get("resp_sst_cargo") or "").strip() or "Responsable SST",
        "vigia_nombre": (empresa.get("vigia_nombre") or "").strip() or "________________",
        "elaboro": (empresa.get("resp_sst_nombre") or "").strip(),
        "aprobo": (empresa.get("rep_legal_nombre") or "").strip(),
        "reviso": (empresa.get("vigia_nombre") or "").strip() or "________________",
    }


def _reemplazar_texto(texto: str, ctx: dict[str, str]) -> str:
    if not texto or not isinstance(texto, str):
        return texto
    out = texto
    razon = _ctx_val(ctx, "razon_social")
    nit = _ctx_val(ctx, "nit")
    arl = _ctx_val(ctx, "arl", fallback="")
    resp = _ctx_val(ctx, "resp_sst_nombre", fallback="")
    rep = _ctx_val(ctx, "rep_legal_nombre", fallback="")
    vigia = _ctx_val(ctx, "vigia_nombre")
    codigo = _ctx_val(ctx, "codigo", fallback="")
    version = _ctx_val(ctx, "version", fallback="")
    fecha = _ctx_val(ctx, "fecha", fallback="")
    ciiu = _ctx_val(ctx, "ciiu", fallback="")

    for k, v in ctx.items():
        val = (v or "").strip() or "________________"
        out = out.replace("{{" + k + "}}", val)
        out = out.replace("{{ " + k + " }}", val)
        out = out.replace("[" + k.upper() + "]", val)

    for rx in _EMPRESA_EJEMPLO_REGEX:
        out = rx.sub(razon, out)
    for rx in _NIT_EJEMPLO_REGEX:
        out = rx.sub(nit, out)

    out = re.sub(
        r"\(CIIU\s*0230\s*y\s*0163\)",
        f"(CIIU {ciiu})" if ciiu else "(CIIU segun actividad de la empresa)",
        out,
        flags=re.I,
    )
    out = re.sub(r"\bCIIU\s*0230\b", f"CIIU {ctx.get('ciiu_codigo') or '____'}", out, flags=re.I)
    out = re.sub(r"\bCIIU\s*0163\b", "", out, flags=re.I)

    def _pref(pat: str, valor: str) -> None:
        nonlocal out
        out = re.sub(pat, lambda m: m.group(1) + valor, out)

    _pref(r"(?i)(raz[o\u00f3]n\s+social\s*:?\s*)_{3,}", razon)
    _pref(r"(?i)(nit\s*:?\s*)_{3,}", nit)
    _pref(r"(?i)(empresa\s*:?\s*)_{3,}", razon)
    _pref(r"(?i)(c[o\u00f3]digo\s*:?\s*)_{3,}", codigo)
    _pref(r"(?i)(versi[o\u00f3]n\s*:?\s*)_{3,}", version)
    _pref(r"(?i)(fecha\s*:?\s*)_{3,}", fecha)

    simples = [
        (r"(?i)\[NOMBRE DE LA EMPRESA\]", razon),
        (r"(?i)\[RAZON SOCIAL\]", razon),
        (r"(?i)\[NIT\]", nit),
        (r"(?i)\[CODIGO\]", codigo),
        (r"(?i)\[VERSION\]", version),
        (r"(?i)\[FECHA\]", fecha),
        (r"(?i)\[ARL\]", arl or "________________"),
        (r"(?i)\[RESPONSABLE SST\]", resp or "________________"),
        (r"(?i)\[REPRESENTANTE LEGAL\]", rep or "________________"),
        (r"(?i)\[VIGIA SST\]", vigia),
        (r"(?i)\[NOMBRE DEL VIG[I\u00cd]A SST\]", vigia),
        (r"(?i)NOMBRE DE LA EMPRESA", razon),
        (r"(?i)Nombre de la empresa", razon),
    ]
    for pat, rep_txt in simples:
        out = re.sub(pat, lambda m, r=rep_txt: r, out)

    out = re.sub(
        r"(?i)(NIT\s*:?\s*)901[\.\s]?989[\.\s]?693\s*-?\s*7",
        lambda m: m.group(1) + nit,
        out,
    )
    return out


def _logo_path(empresa: dict) -> Path | None:
    p = (empresa.get("logo_path") or "").strip()
    if not p:
        return None
    path = Path(p)
    if not path.exists() or path.name == "logo.png":
        return None
    return path


def _set_paragraph_text(p, nuevo: str) -> None:
    if p.runs:
        p.runs[0].text = nuevo
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(nuevo)


def _iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_all_paragraphs(doc) -> Iterable:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                yield p
            for table in part.tables:
                yield from _iter_table_paragraphs(table)


def _aplicar_documento_word(doc, ctx: dict[str, str]) -> None:
    for p in _iter_all_paragraphs(doc):
        if not p.text:
            continue
        nuevo = _reemplazar_texto(p.text, ctx)
        if nuevo != p.text:
            _set_paragraph_text(p, nuevo)


def _flex_phrase_pattern(phrase: str) -> str:
    """Permite etiquetas XML entre palabras (texto partido en runs de Word)."""
    gap = r"(?:\s|&nbsp;|&#160;|<[^>]+>)*"
    parts = [re.escape(p) for p in phrase.split() if p]
    return gap.join(parts)


def _ooxml_scrub_patterns(ctx: dict[str, str]) -> list[tuple[re.Pattern[str], str]]:
    razon = xml_escape(_ctx_val(ctx, "razon_social"))
    nit = xml_escape(_ctx_val(ctx, "nit"))
    phrases = [
        "FOREST GREEN SERVICIOS AMBIENTALES S.A.S.",
        "FOREST GREEN SERVICIOS AMBIENTALES S.A.S",
        "FOREST GREEN SERVICIOS AMBIENTALES",
        "FOREST GREEN",
        "Laboratorios LT S.A.S.",
        "Laboratorios LT S.A.S",
        "XXXXXX XXXXXX",
        "NOMBRE DE LA EMPRESA",
    ]
    pats: list[tuple[re.Pattern[str], str]] = []
    for ph in phrases:
        pats.append((re.compile(_flex_phrase_pattern(ph), re.I), razon))
    # NIT demo con posibles tags entre digitos
    nit_flex = (
        r"901"
        + r"(?:\s|<[^>]+>)*"
        + r"[\.]?"
        + r"(?:\s|<[^>]+>)*"
        + r"989"
        + r"(?:\s|<[^>]+>)*"
        + r"[\.]?"
        + r"(?:\s|<[^>]+>)*"
        + r"693"
        + r"(?:\s|<[^>]+>)*"
        + r"-?"
        + r"(?:\s|<[^>]+>)*"
        + r"7"
    )
    pats.append((re.compile(nit_flex, re.I), nit))
    return pats


def _logo_bytes_for_ext(logo_path: Path, ext: str) -> bytes | None:
    """Convierte el logo de la empresa al formato del media embebido (jpg/png)."""
    try:
        from PIL import Image
    except Exception:
        return logo_path.read_bytes() if logo_path.suffix.lower() == f".{ext}" else None
    try:
        img = Image.open(logo_path)
        if img.mode in ("RGBA", "LA", "P"):
            background = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == "P":
                img = img.convert("RGBA")
            background.paste(img, mask=img.split()[-1] if img.mode == "RGBA" else None)
            img = background
        elif img.mode != "RGB":
            img = img.convert("RGB")
        # Limitar tamano tipico de encabezado
        img.thumbnail((220, 120))
        bio = io.BytesIO()
        ext_l = ext.lower().lstrip(".")
        if ext_l in {"jpg", "jpeg"}:
            img.save(bio, format="JPEG", quality=90)
        elif ext_l == "png":
            img.save(bio, format="PNG")
        else:
            img.save(bio, format="PNG")
        return bio.getvalue()
    except Exception:
        try:
            return logo_path.read_bytes()
        except Exception:
            return None


def _blank_logo_bytes(ext: str) -> bytes:
    """Imagen blanca pequena para quitar logos de ejemplo si no hay logo propio."""
    from PIL import Image

    img = Image.new("RGB", (160, 80), (255, 255, 255))
    bio = io.BytesIO()
    ext_l = ext.lower().lstrip(".")
    if ext_l in {"jpg", "jpeg"}:
        img.save(bio, format="JPEG", quality=85)
    else:
        img.save(bio, format="PNG")
    return bio.getvalue()


def _ooxml_header_meta_patterns(ctx: dict[str, str]) -> list[tuple[re.Pattern[str], str]]:
    """Fecha / Version / Codigo FG-FOR-SST del encabezado de plantilla."""
    fecha = xml_escape(_ctx_val(ctx, "fecha"))
    version = xml_escape(_ctx_val(ctx, "version"))
    codigo = xml_escape(_ctx_val(ctx, "codigo"))
    gap = r"(?:\s|<[^>]+>)*"
    pats: list[tuple[re.Pattern[str], str]] = []
    # Fecha: 01/02/2026  (con posibles tags XML entre tokens)
    pats.append(
        (
            re.compile(
                rf"(Fecha{gap}:{gap})\d{{1,2}}{gap}/{gap}\d{{1,2}}{gap}/{gap}\d{{2,4}}",
                re.I,
            ),
            rf"\g<1>{fecha}",
        )
    )
    pats.append(
        (
            re.compile(
                rf"(Versi(?:\u00f3|o)n{gap}:{gap})\d{{1,4}}",
                re.I,
            ),
            rf"\g<1>{version}",
        )
    )
    pats.append(
        (
            re.compile(
                rf"(C\u00f3digo{gap}:{gap}|Codigo{gap}:{gap})FG{gap}-{gap}FOR{gap}-{gap}SST{gap}-{gap}\d+",
                re.I,
            ),
            rf"\g<1>{codigo}",
        )
    )
    # Codigo suelto FG-FOR-SST-39
    pats.append(
        (
            re.compile(
                rf"FG{gap}-{gap}FOR{gap}-{gap}SST{gap}-{gap}\d+",
                re.I,
            ),
            codigo,
        )
    )
    return pats


def _resolve_media_target(target: str, names: set[str], rel_name: str) -> str | None:
    """Normaliza Target de .rels (Word/Excel) a una ruta real del zip."""
    t = (target or "").replace("\\", "/").strip()
    if not t:
        return None
    candidates = [
        t.lstrip("/"),
        t,
    ]
    # Relativo al archivo .rels (p.ej. ../media/image1.jpeg desde word/_rels/)
    if t.startswith("../") or t.startswith("./") or not t.startswith("/") and "/" not in t.split("/")[0]:
        base = str(Path(rel_name).parent).replace("\\", "/")
        # parent of _rels is word/ or xl/drawings/
        if base.endswith("/_rels"):
            base = base[: -len("/_rels")]
        joined = str((Path(base) / t).as_posix())
        # normalize .. 
        parts: list[str] = []
        for part in joined.split("/"):
            if part in ("", "."):
                continue
            if part == "..":
                if parts:
                    parts.pop()
                continue
            parts.append(part)
        candidates.append("/".join(parts))
    if "media/" in t.lower():
        # Excel a veces usa /xl/media/... o media/image1.jpeg
        if t.lstrip("/").startswith("xl/media/") or t.lstrip("/").startswith("word/media/"):
            candidates.append(t.lstrip("/"))
        elif "xl/" in rel_name.replace("\\", "/").lower():
            candidates.append("xl/" + t.lstrip("/").removeprefix("xl/"))
            if t.lower().startswith("media/") or "/media/" in t.lower():
                candidates.append("xl/media/" + Path(t).name)
        elif "word/" in rel_name.replace("\\", "/").lower():
            candidates.append("word/media/" + Path(t).name)
    for cand in candidates:
        c = cand.lstrip("/")
        if c in names:
            return c
    return None


def _scrub_ooxml_file(path: Path, ctx: dict[str, str], logo_path: Path | None = None) -> None:
    """Reemplazo en XML + sustitucion de logos embebidos (Word y Excel)."""
    if not path.exists():
        return
    pats = _ooxml_scrub_patterns(ctx) + _ooxml_header_meta_patterns(ctx)

    logo_media_targets: set[str] = set()
    with zipfile.ZipFile(path, "r") as zin:
        names = set(zin.namelist())
        for rel_name in list(names):
            low = rel_name.replace("\\", "/").lower()
            if not low.endswith(".rels"):
                continue
            # Word headers/footers + Excel drawings (logos en planillas)
            useful = (
                "/_rels/header" in low
                or "/_rels/footer" in low
                or "xl/drawings/_rels/" in low
                or low.endswith("drawing1.xml.rels")
                or "drawings/_rels/" in low
            )
            if not useful:
                continue
            try:
                rel_xml = zin.read(rel_name).decode("utf-8")
            except Exception:
                continue
            for m in re.finditer(
                r'Type="[^"]*image"[^>]*Target="([^"]+)"|Target="([^"]*media/[^"]+)"[^>]*Type="[^"]*image"',
                rel_xml,
                re.I,
            ):
                target = m.group(1) or m.group(2)
                resolved = _resolve_media_target(target, names, rel_name)
                if resolved:
                    logo_media_targets.add(resolved)
            # Fallback mas simple: cualquier Target a media/
            for m in re.finditer(r'Target="([^"]*media/[^"]+)"', rel_xml, re.I):
                resolved = _resolve_media_target(m.group(1), names, rel_name)
                if resolved and resolved.lower().endswith((".jpg", ".jpeg", ".png")):
                    logo_media_targets.add(resolved)

        medias = [
            n
            for n in names
            if "/media/" in n.lower() and n.lower().endswith((".jpg", ".jpeg", ".png"))
        ]
        # Excel: image1 suele ser el logo FOREST GREEN
        if not logo_media_targets:
            for n in medias:
                if re.search(r"image1\.(jpg|jpeg|png)$", n, re.I):
                    logo_media_targets.add(n)
        if not logo_media_targets and len(medias) == 1:
            logo_media_targets.add(medias[0])
        # Si hay varias en xl/media y ninguna elegida, reemplazar la primera (logo)
        if not logo_media_targets and medias:
            xl_medias = sorted(n for n in medias if n.lower().startswith("xl/media/"))
            if xl_medias:
                logo_media_targets.add(xl_medias[0])

    buf = io.BytesIO()
    changed = False
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(
        buf, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            name = info.filename
            name_l = name.lower()

            if name in logo_media_targets or name.replace("\\", "/") in logo_media_targets:
                # Solo reemplazar si hay logo de empresa; no blanquear el de la plantilla
                ext = Path(name).suffix.lstrip(".").lower()
                if logo_path and logo_path.exists():
                    nuevo_bytes = _logo_bytes_for_ext(logo_path, ext)
                    if nuevo_bytes is not None:
                        data = nuevo_bytes
                        changed = True

            elif name_l.endswith((".xml", ".rels")) and not name_l.endswith(".bin"):
                try:
                    text = data.decode("utf-8")
                except UnicodeDecodeError:
                    zout.writestr(info, data)
                    continue
                nuevo = text
                for rx, rep in pats:
                    try:
                        nuevo2 = rx.sub(rep, nuevo)
                    except re.error:
                        continue
                    if nuevo2 != nuevo:
                        changed = True
                        nuevo = nuevo2
                data = nuevo.encode("utf-8")
            zout.writestr(info, data)
    if changed:
        path.write_bytes(buf.getvalue())


def rellenar_word(
    plantilla: Path,
    empresa: dict,
    doc_meta: dict,
    version: int,
    fecha: str,
    out_path: Path,
) -> Path:
    ctx = _contexto(empresa, doc_meta, version, fecha)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    rendered = False
    try:
        from docxtpl import DocxTemplate

        tpl = DocxTemplate(str(plantilla))
        context: dict[str, Any] = {k: _ctx_val(ctx, k) for k in ctx}
        logo = _logo_path(empresa)
        if logo:
            try:
                from docxtpl import InlineImage
                from docx.shared import Mm

                context["logo"] = InlineImage(tpl, str(logo), width=Mm(25))
            except Exception:
                pass
        tpl.render(context)
        tpl.save(str(out_path))
        rendered = True
    except Exception:
        rendered = False

    from docx import Document

    if not rendered:
        doc = Document(str(plantilla))
        doc.save(str(out_path))

    doc = Document(str(out_path))
    _aplicar_documento_word(doc, ctx)
    doc.save(str(out_path))
    logo = _logo_path(empresa)
    _scrub_ooxml_file(out_path, ctx, logo_path=logo)
    # Reparar layouts rotos de plantillas (encuestas con columnas fingidas)
    try:
        from generador.layout_encuestas import reparar_layouts_word

        reparar_layouts_word(out_path, fecha=ctx.get("fecha"))
        # Reaplicar scrub de logo/meta por si el repair reescribio el docx
        _scrub_ooxml_file(out_path, ctx, logo_path=logo)
    except Exception:
        pass
    # Siempre inyectar logo + razon social + NIT en el encabezado
    try:
        from generador.encabezado_empresa import inyectar_encabezado_empresa_word

        inyectar_encabezado_empresa_word(
            out_path, empresa, doc_meta, version, fecha
        )
    except Exception:
        pass
    return out_path



def _cell_fill_hex(cell) -> str:
    fill = getattr(cell, "fill", None)
    if not fill or getattr(fill, "fill_type", None) != "solid":
        return ""
    color = getattr(fill, "fgColor", None)
    if not color:
        return ""
    # Ignorar theme/indexed (no son hex reales)
    if getattr(color, "type", None) not in ("rgb", None) and getattr(color, "type", None) != "rgb":
        t = str(getattr(color, "type", ""))
        if t and t != "rgb":
            return ""
    rgb = getattr(color, "rgb", None)
    if not rgb or not isinstance(rgb, str):
        return ""
    hx = rgb[-6:].upper()
    if len(hx) != 6 or any(c not in "0123456789ABCDEF" for c in hx):
        return ""
    return hx


def _aplicar_colores_excel(wb, empresa: dict) -> None:
    from openpyxl.styles import Font, PatternFill
    from generador.colores_marca import (
        colores_desde_logo,
        elegir_tono_marca,
        es_relleno_marca_plantilla,
    )

    logo = _logo_path(empresa)
    palette = colores_desde_logo(logo)
    fill_primary = PatternFill("solid", fgColor=palette["primary"])
    fill_secondary = PatternFill("solid", fgColor=palette["secondary"])
    fill_light = PatternFill("solid", fgColor=palette["light"])
    on_pri = palette.get("on_primary", "FFFFFF")
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                hx = _cell_fill_hex(cell)
                if not hx or not es_relleno_marca_plantilla(hx):
                    continue
                nuevo = elegir_tono_marca(hx, palette)
                if nuevo == palette["light"]:
                    cell.fill = fill_light
                elif nuevo == palette["secondary"]:
                    cell.fill = fill_secondary
                else:
                    cell.fill = fill_primary
                if nuevo in (palette["primary"], palette["secondary"]):
                    prev = cell.font
                    cell.font = Font(
                        name=prev.name,
                        size=prev.size,
                        bold=prev.bold,
                        italic=prev.italic,
                        color=on_pri,
                    )


_LOGO_PLACEHOLDER_RE = re.compile(
    r"^\s*(?:\{\{\s*)?logo(?:\s*\}\})?\s*$",
    re.I,
)


def _es_placeholder_logo(valor: Any) -> bool:
    return isinstance(valor, str) and bool(_LOGO_PLACEHOLDER_RE.match(valor.strip()))


def _rellenar_bloque_meta_excel(texto: str, ctx: dict[str, str]) -> str:
    """Completa bloques tipo CODIGO/VERSION/FECHA e inserta empresa/NIT si faltan."""
    if not texto or not isinstance(texto, str):
        return texto
    out = texto
    razon = (ctx.get("razon_social") or "").strip()
    nit = (ctx.get("nit") or "").strip()
    codigo = (ctx.get("codigo") or "").strip()
    version = (ctx.get("version") or "").strip()
    fecha = (ctx.get("fecha") or "").strip()

    def _fill_line(pat: str, valor: str) -> None:
        nonlocal out
        if not valor:
            return
        out = re.sub(
            pat,
            lambda m: m.group(1) + valor,
            out,
            count=1,
            flags=re.I | re.M,
        )

    _fill_line(r"(c[o\u00f3]digo\s*:\s*)\s*$", codigo)
    _fill_line(r"(c[o\u00f3]digo\s*:\s*)(?=\n|$)", codigo)
    _fill_line(r"(versi[o\u00f3]n\s*:\s*)\s*(?:0?1)?\s*(?=\n|$)", version)
    _fill_line(r"(fecha\s*:\s*)\s*(?=\n|$)", fecha)
    _fill_line(r"(fecha\s+de\s+actualizaci[o\u00f3]n\s*:\s*)\s*$", fecha)

    looks_meta = bool(
        re.search(r"(?i)c[o\u00f3]digo\s*:", out)
        and re.search(r"(?i)versi[o\u00f3]n\s*:", out)
    )
    if looks_meta and razon and not re.search(r"(?i)\bempresa\s*:", out):
        encabezado = f"EMPRESA: {razon}"
        if nit:
            encabezado += f"\nNIT: {nit}"
        out = encabezado + "\n" + out.lstrip()
    return out


def _inyectar_logo_y_empresa_excel(wb, empresa: dict, ctx: dict[str, str]) -> None:
    """Plantillas nuevas a veces traen texto 'logo' sin imagen embebida ni datos de empresa."""
    logo = _logo_path(empresa)
    razon = (ctx.get("razon_social") or "").strip()
    nit = (ctx.get("nit") or "").strip()
    resp = (ctx.get("resp_sst_nombre") or "").strip()
    fecha = (ctx.get("fecha") or "").strip()

    for ws in wb.worksheets:
        anclas_logo: list[str] = []
        max_col = min(ws.max_column or 1, 80)
        max_row_scan = min(ws.max_row or 1, 12)

        for row in ws.iter_rows(min_row=1, max_row=max_row_scan, max_col=max_col):
            for cell in row:
                if not isinstance(cell.value, str):
                    continue
                raw = cell.value
                if _es_placeholder_logo(raw):
                    anclas_logo.append(cell.coordinate)
                    # Dejar celda limpia; la imagen va anclada ahi
                    cell.value = None if logo else (
                        f"{razon}\nNIT: {nit}" if razon else None
                    )
                    continue

                nuevo = _rellenar_bloque_meta_excel(raw, ctx)
                # Elaborado por generico de plantilla
                if resp and re.search(
                    r"(?i)^elaborado\s+por\s*:\s*(encargado\s+del\s+sg-?sst.?|responsable\s+sst)?\s*$",
                    nuevo.strip(),
                ):
                    nuevo = f"Elaborado por: {resp}"
                elif resp and re.search(r"(?i)^elaborado\s+por\s*:\s*$", nuevo.strip()):
                    nuevo = f"Elaborado por: {resp}"

                if fecha and re.search(
                    r"(?i)^fecha\s+de\s+actualizaci[o\u00f3]n\s*:\s*$",
                    nuevo.strip(),
                ):
                    nuevo = f"Fecha de actualizaci\u00f3n: {fecha}"

                if nuevo != raw:
                    cell.value = nuevo

        # Si no habia placeholder pero tampoco imagen, poner datos en A1 si esta vacia
        if not anclas_logo and not getattr(ws, "_images", None):
            a1 = ws.cell(1, 1)
            if a1.value is None and razon:
                a1.value = f"{razon}\nNIT: {nit}" if nit else razon
                anclas_logo.append("A1")

        if logo and anclas_logo:
            try:
                img = XLImage(str(logo))
                img.width = 96
                img.height = 58
                ws.add_image(img, anclas_logo[0])
            except Exception:
                # Si falla la imagen, al menos deja texto de empresa
                cell = ws[anclas_logo[0]]
                if cell.value is None and razon:
                    cell.value = f"{razon}\nNIT: {nit}" if nit else razon

        # Garantizar empresa visible en encabezado si sigue ausente
        header_txt = " ".join(
            str(ws.cell(r, c).value or "")
            for r in range(1, min(4, max_row_scan + 1))
            for c in range(1, min(max_col, 70) + 1)
        )
        if razon and razon.lower() not in header_txt.lower():
            # Preferir celda meta (CODIGO/VERSION); si no, A1
            colocado = False
            for row in ws.iter_rows(min_row=1, max_row=3, max_col=max_col):
                for cell in row:
                    val = cell.value
                    if isinstance(val, str) and re.search(r"(?i)c[o\u00f3]digo\s*:", val):
                        pref = f"EMPRESA: {razon}"
                        if nit:
                            pref += f"\nNIT: {nit}"
                        cell.value = pref + "\n" + val.lstrip()
                        colocado = True
                        break
                if colocado:
                    break
            if not colocado:
                a1 = ws.cell(1, 1)
                if a1.value is None:
                    a1.value = f"{razon}\nNIT: {nit}" if nit else razon


def asegurar_marca_excel(
    out_path: Path | str,
    empresa: dict,
    doc_meta: dict,
    version: int | str,
    fecha: str,
) -> None:
    """Pase final: garantiza logo + empresa en cualquier Excel generado."""
    path = Path(out_path)
    if not path.exists():
        return
    ver_int = int(version) if str(version).isdigit() else 1
    ctx = _contexto(empresa, doc_meta, ver_int, str(fecha))
    ctx["codigo"] = str(doc_meta.get("codigo") or ctx.get("codigo") or "")
    ctx["version"] = str(version)
    ctx["fecha"] = str(fecha)
    ctx["razon_social"] = (empresa.get("razon_social") or "").strip()
    ctx["nit"] = (empresa.get("nit") or "").strip()
    wb = load_workbook(path)
    _inyectar_logo_y_empresa_excel(wb, empresa, ctx)
    ws = wb.worksheets[0]
    razon = ctx["razon_social"]
    nit = ctx["nit"]
    header_blob = " ".join(
        str(ws.cell(r, c).value or "")
        for r in range(1, 5)
        for c in range(1, min(ws.max_column or 1, 80) + 1)
    )
    if razon and razon.lower() not in header_blob.lower():
        ws.insert_rows(1)
        ws.cell(
            1,
            1,
            f"EMPRESA: {razon} | NIT: {nit} | Codigo: {ctx['codigo']} | v{version} | {fecha}",
        )
    logo = _logo_path(empresa)
    if logo and not getattr(ws, "_images", None):
        try:
            img = XLImage(str(logo))
            img.width = 96
            img.height = 58
            ws.add_image(img, "A1")
        except Exception:
            pass
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    _scrub_ooxml_file(path, ctx, logo_path=logo)


def rellenar_excel(
    plantilla: Path,
    empresa: dict,
    doc_meta: dict,
    version: int,
    fecha: str,
    out_path: Path,
) -> Path:
    ctx = _contexto(empresa, doc_meta, version, fecha)
    wb = load_workbook(plantilla)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    nuevo = _reemplazar_texto(cell.value, ctx)
                    if nuevo != cell.value:
                        cell.value = nuevo
    try:
        _inyectar_logo_y_empresa_excel(wb, empresa, ctx)
    except Exception:
        pass
    try:
        _aplicar_colores_excel(wb, empresa)
    except Exception:
        pass
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    logo = _logo_path(empresa)
    _scrub_ooxml_file(out_path, ctx, logo_path=logo)
    try:
        asegurar_marca_excel(out_path, empresa, doc_meta, version, fecha)
    except Exception:
        pass
    return out_path

