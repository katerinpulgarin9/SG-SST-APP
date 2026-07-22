# -*- coding: utf-8 -*-
"""Motor de generacion: plantilla si existe, si no genera desde cero."""
from __future__ import annotations

import re
import tempfile
from datetime import date
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter

from generador import db
from generador.catalogo_maestro import codigo_sst
from generador.plantillas_index import (
    encontrar_plantilla,
    estado_plantillas_catalogo,
    listar_plantillas,
)
from generador.relleno_plantilla import (
    rellenar_excel,
    rellenar_word,
    _scrub_ooxml_file,
    _contexto,
    _logo_path,
    asegurar_marca_excel,
)
from generador.encabezado_empresa import inyectar_encabezado_empresa_word
from generador.colores_marca import colores_desde_logo
from generador.familias_documento import (
    COLUMNAS_EXCEL,
    contexto_actividad,
    peligros_sugeridos,
    secciones_familia,
)

BLACK = RGBColor(0, 0, 0)
THIN = Side(style="thin", color="000000")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
HEADER_FILL = PatternFill("solid", fgColor="D9E1F2")

# Matching de plantillas: generador.plantillas_index


def _slug(nombre: str) -> str:
    s = re.sub(r"[^\w\s-]", "", nombre, flags=re.UNICODE)
    s = re.sub(r"\s+", "_", s.strip())
    return s[:60] or "Documento"


def _force_black(run) -> None:
    run.font.color.rgb = BLACK
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:color", "w:themeColor"):
        existing = r_pr.find(qn(tag))
        if existing is not None:
            r_pr.remove(existing)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)


def _table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def _logo_path(empresa: dict) -> Optional[Path]:
    """Solo logo adjunto por empresa; sin fallback al logo por defecto."""
    raw = (empresa.get("logo_path") or "").strip()
    if not raw:
        return None
    p = Path(raw)
    if not p.exists():
        return None
    # Evitar reutilizar logo.png global de otra empresa
    if p.name == "logo.png":
        return None
    return p


def _shade_word_cell(cell, hex_color: str) -> None:
    """Relleno de celda Word (hex RRGGBB)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_word_header(doc: Document, empresa: dict, titulo: str, codigo: str, version: int, fecha: str) -> None:
    palette = colores_desde_logo(_logo_path(empresa))
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(table)
    for i, w in enumerate([Inches(1.6), Inches(3.8), Inches(2.2)]):
        table.columns[i].width = w

    logo_cell = table.cell(0, 0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _shade_word_cell(logo_cell, palette["light"])
    p = logo_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = _logo_path(empresa)
    if logo:
        p.add_run().add_picture(str(logo), width=Inches(1.0))

    title_cell = table.cell(0, 1)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _shade_word_cell(title_cell, palette["primary"])
    pt = title_cell.paragraphs[0]
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pt.add_run(titulo.upper())
    r.bold = True
    r.font.size = Pt(12)
    try:
        r.font.color.rgb = RGBColor.from_string(palette["on_primary"])
    except Exception:
        _force_black(r)

    meta_cell = table.cell(0, 2)
    meta_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _shade_word_cell(meta_cell, palette["secondary"])
    meta = meta_cell.paragraphs[0]
    mr = meta.add_run(
        f"{empresa.get('razon_social','')}\nNIT: {empresa.get('nit','')}\n"
        f"Codigo: {codigo}\nVersion: {version}\nFecha: {fecha}"
    )
    mr.font.size = Pt(9)
    try:
        mr.font.color.rgb = RGBColor.from_string(palette["on_primary"])
    except Exception:
        _force_black(mr)
    doc.add_paragraph()



def add_word_signatures(doc: Document, empresa: dict, include_vigia: bool = True) -> None:
    """Bloque compacto de firmas al final (no debe dominar la primera pagina)."""
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("FIRMAS")
    r_title.bold = True
    r_title.font.size = Pt(11)
    _force_black(r_title)

    cols = 3 if include_vigia else 2
    table = doc.add_table(rows=4, cols=cols)
    _table_borders(table)

    headers = ["Elabor\u00f3", "Aprob\u00f3"]
    names = [
        empresa.get("resp_sst_nombre") or "________________",
        empresa.get("rep_legal_nombre") or "________________",
    ]
    cargos = [
        empresa.get("resp_sst_cargo") or "Responsable SST",
        empresa.get("rep_legal_cargo") or "Representante legal",
    ]
    if include_vigia:
        headers.append("Revis\u00f3")
        vigia = (empresa.get("vigia_nombre") or "").strip()
        names.append(vigia if vigia else "________________")
        cargos.append("Vig\u00eda en SST")

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        _force_black(run)

    rows = [
        [f"Nombre: {n}" for n in names],
        [f"Cargo: {c}" for c in cargos],
        ["Firma / Fecha: ______________"] * cols,
    ]
    for ri, vals in enumerate(rows, start=1):
        for ci, val in enumerate(vals):
            cell = table.cell(ri, ci)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8)
            _force_black(run)


def _include_vigia(doc_meta: dict) -> bool:
    fam = (doc_meta.get("familia") or "").lower()
    nombre = (doc_meta.get("nombre") or "").lower()
    if fam in {"matriz", "plan", "formato", "programa", "evaluacion"}:
        return True
    keys = ("emergencia", "epp", "inspeccion", "inspecci\u00f3n", "peligro", "riesgo")
    return any(k in nombre for k in keys)


def _contar_parrafos_utiles(doc: Document) -> int:
    """Cuenta parrafos con texto sustancial (cuerpo, no solo titulos cortos)."""
    n = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if len(t) >= 40:
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = (p.text or "").strip()
                    if len(t) >= 40:
                        n += 1
    return n


def _agregar_secciones_word(document: Document, empresa: dict, doc_meta: dict) -> None:
    for heading, paragraphs in secciones_familia(
        doc_meta.get("familia") or "", empresa, doc_meta.get("nombre") or ""
    ):
        p = document.add_paragraph()
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(12)
        _force_black(run)
        for text in paragraphs:
            document.add_paragraph(text)


MIN_PARRAFOS_PLANTILLA = 15


def generar_word_desde_cero(empresa: dict, doc_meta: dict, version: int, fecha: str, out_path: Path) -> Path:
    document = Document()
    add_word_header(document, empresa, doc_meta["nombre"], doc_meta["codigo"], version, fecha)
    _agregar_secciones_word(document, empresa, doc_meta)
    add_word_signatures(document, empresa, include_vigia=_include_vigia(doc_meta))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    try:
        inyectar_encabezado_empresa_word(
            out_path, empresa, doc_meta, version, fecha
        )
    except Exception:
        pass
    return out_path


def generar_word_desde_plantilla(empresa: dict, doc_meta: dict, version: int, fecha: str, plantilla: Path, out_path: Path) -> Path:
    """Usa la plantilla como base; si queda corta, agrega contenido ampliado."""
    try:
        rellenar_word(plantilla, empresa, doc_meta, version, fecha, out_path)
    except Exception:
        return generar_word_desde_cero(empresa, doc_meta, version, fecha, out_path)

    doc = Document(str(out_path))
    if _contar_parrafos_utiles(doc) < MIN_PARRAFOS_PLANTILLA:
        sep = doc.add_paragraph()
        run = sep.add_run(
            "CONTENIDO COMPLEMENTARIO DEL SG-SST "
            "(la plantilla base era breve; se ampl\u00eda para uso operativo)"
        )
        run.bold = True
        run.font.size = Pt(11)
        _force_black(run)
        _agregar_secciones_word(doc, empresa, doc_meta)
        add_word_signatures(doc, empresa, include_vigia=_include_vigia(doc_meta))
        doc.save(str(out_path))
    # Segunda pasada anti FOREST GREEN (XML partido / encabezados)
    try:
        _scrub_ooxml_file(
            out_path,
            _contexto(empresa, doc_meta, version, fecha),
            logo_path=_logo_path(empresa),
        )
    except Exception:
        pass
    # Forzar logo + empresa en encabezado (plantillas no traen razon social/NIT)
    try:
        inyectar_encabezado_empresa_word(
            out_path, empresa, doc_meta, version, fecha
        )
    except Exception:
        pass
    return out_path


def _excel_header(ws, empresa: dict, titulo: str, codigo: str, version: int, fecha: str) -> int:
    pal = colores_desde_logo(_logo_path(empresa))
    fill_pri = PatternFill("solid", fgColor=pal["primary"])
    fill_sec = PatternFill("solid", fgColor=pal["secondary"])
    fill_light = PatternFill("solid", fgColor=pal["light"])
    on_pri = pal.get("on_primary", "FFFFFF")
    ws.merge_cells("A1:B3")
    ws.merge_cells("C1:H3")
    ws.merge_cells("I1:L3")
    logo = _logo_path(empresa)
    if logo:
        try:
            img = XLImage(str(logo))
            img.width = 90
            img.height = 55
            ws.add_image(img, "A1")
        except Exception:
            pass
    c = ws.cell(1, 3, titulo.upper())
    c.font = Font(bold=True, size=12, color=on_pri)
    c.fill = fill_pri
    c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    meta = ws.cell(
        1, 9,
        f"{empresa.get('razon_social','')}\nNIT: {empresa.get('nit','')}\n"
        f"C?digo: {codigo}\nVersi?n: {version}\nFecha: {fecha}",
    )
    meta.font = Font(size=9, color=on_pri)
    meta.fill = fill_sec
    meta.alignment = Alignment(vertical="center", wrap_text=True)
    for r in range(1, 4):
        for col in range(1, 13):
            cell = ws.cell(r, col)
            cell.border = BORDER
            if col <= 2:
                cell.fill = fill_light
            elif 3 <= col <= 8:
                cell.fill = fill_pri
            elif col >= 9:
                cell.fill = fill_sec
    # Restaurar texto titulo/meta (merge usa celda origen)
    c = ws.cell(1, 3)
    c.font = Font(bold=True, size=12, color=on_pri)
    c.fill = fill_pri
    meta = ws.cell(1, 9)
    meta.font = Font(size=9, color=on_pri)
    meta.fill = fill_sec
    return 5


def _excel_signatures(ws, row: int, empresa: dict) -> None:
    blocks = [
        ("Elabor?", empresa.get("resp_sst_nombre") or "________________", empresa.get("resp_sst_cargo") or "Responsable SST"),
        ("Aprob?", empresa.get("rep_legal_nombre") or "________________", empresa.get("rep_legal_cargo") or "Representante legal"),
        ("Revis?", (empresa.get("vigia_nombre") or "").strip() or "________________", "Vig?a en SST"),
    ]
    span = 4
    total = span * len(blocks)
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=total)
    cell = ws.cell(row, 1, "FIRMAS")
    cell.font = Font(bold=True, color="000000")
    for c in range(1, total + 1):
        ws.cell(row, c).border = BORDER
    row += 1
    for i, (label, nombre, cargo) in enumerate(blocks):
        c0 = i * span + 1
        for offset, txt in enumerate(
            [label, f"Nombre: {nombre}", f"Cargo: {cargo}", "Firma: _______________________", "Fecha: _______________________"]
        ):
            rr = row + offset
            ws.merge_cells(start_row=rr, start_column=c0, end_row=rr, end_column=c0 + span - 1)
            cell = ws.cell(rr, c0, txt)
            cell.font = Font(bold=(offset == 0), size=10, color="000000")
            for c in range(c0, c0 + span):
                ws.cell(rr, c).border = BORDER


def generar_excel_desde_cero(empresa: dict, doc_meta: dict, version: int, fecha: str, out_path: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Documento"
    r = _excel_header(ws, empresa, doc_meta["nombre"], doc_meta["codigo"], version, fecha)
    ws.cell(r, 1, f"Empresa: {empresa.get('razon_social','')} | NIT: {empresa.get('nit','')} | {contexto_actividad(empresa)}")
    r += 2
    cols = COLUMNAS_EXCEL.get(doc_meta["familia"], COLUMNAS_EXCEL["formato"])
    for c, h in enumerate(cols, 1):
        cell = ws.cell(r, c, h)
        cell.font = Font(bold=True, color="000000")
        _pal = colores_desde_logo(_logo_path(empresa))
        cell.fill = PatternFill("solid", fgColor=_pal["light"])
        cell.border = BORDER
        cell.alignment = Alignment(wrap_text=True, horizontal="center", vertical="center")
    header_row = r

    # Filas de ejemplo seg?n familia
    ejemplos = []
    if doc_meta["familia"] == "matriz":
        for i, p in enumerate(peligros_sugeridos(empresa)[:8], 1):
            ejemplos.append([i, "Proceso principal", "Actividad operativa", p, "Por valorar", "Controles a definir", empresa.get("resp_sst_nombre", ""), fecha])
    elif doc_meta["familia"] == "evaluacion":
        ejemplos = [
            ["I. PLANEAR", "Recursos", "1.1.1", "Responsable SG-SST designado", 0.5, "", "", "", ""],
            ["I. PLANEAR", "Pol?tica", "2.1.1", "Pol?tica SST divulgada", 1.0, "", "", "", ""],
            ["II. HACER", "IPVR", "12.1.1", "Matriz de peligros actualizada", 2.0, "", "", "", ""],
        ]
    else:
        for i in range(1, 6):
            ejemplos.append([i] + [""] * (len(cols) - 1))

    for row_data in ejemplos:
        r += 1
        for c, val in enumerate(row_data[: len(cols)], 1):
            cell = ws.cell(r, c, val)
            cell.border = BORDER
            cell.alignment = Alignment(wrap_text=True, vertical="center")

    for c in range(1, len(cols) + 1):
        ws.column_dimensions[get_column_letter(c)].width = 16 if c > 1 else 8
    if len(cols) >= 4:
        ws.column_dimensions["D"].width = 36

    _excel_signatures(ws, r + 2, empresa)
    ws.sheet_view.showGridLines = False
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return out_path


def generar_excel_desde_plantilla(empresa: dict, doc_meta: dict, version: int, fecha: str, plantilla: Path, out_path: Path) -> Path:
    """Usa la plantilla Excel como base y rellena datos de empresa."""
    try:
        return rellenar_excel(plantilla, empresa, doc_meta, version, fecha, out_path)
    except Exception:
        return generar_excel_desde_cero(empresa, doc_meta, version, fecha, out_path)



def generar_documento(
    codigo: str,
    fecha: Optional[str] = None,
    version: Optional[int] = None,
    auto_bump: bool = False,
    numero_asignado: Optional[int] = None,
    plantilla_path: Optional[Path | str] = None,
    forzar_sin_plantilla: bool = False,
) -> dict[str, Any]:
    """Genera un documento en memoria (temp) y devuelve bytes para descarga.

    `codigo` identifica el documento en el catalogo.
    `numero_asignado` (opcional) fija el NNN del codigo final SST-SG-[ABR]-[NNN]
    segun el orden de generacion en la app.
    """
    db.init_db()
    empresa = db.get_empresa()
    doc_meta = db.get_documento(codigo)
    if not doc_meta:
        raise ValueError(f"No existe el documento {codigo} en el catalogo")
    if not (empresa.get("razon_social") or "").strip():
        raise ValueError(
            "La empresa activa no tiene razon social. "
            "Ve a Empresa, completa los datos (y el logo) y vuelve a generar."
        )

    fecha = fecha or date.today().strftime("%d/%m/%Y")
    if version is not None:
        version = db.set_version(codigo, version)
    elif auto_bump:
        version = db.bump_version(codigo)
    else:
        actual = int(doc_meta.get("version") or 0)
        version = db.set_version(codigo, actual if actual >= 1 else 1)

    abr = (doc_meta.get("abreviatura") or "FR").upper()
    if numero_asignado is not None:
        n = max(1, int(numero_asignado))
        codigo_final = codigo_sst(abr, n)
        if n > db.get_ultimo_consecutivo():
            db.set_ultimo_consecutivo(n)
    else:
        n = db.next_consecutivo()
        codigo_final = codigo_sst(abr, n)

    doc_gen = dict(doc_meta)
    doc_gen["codigo"] = codigo_final

    if forzar_sin_plantilla:
        plantilla = None
    elif plantilla_path:
        plantilla = Path(plantilla_path)
        if not plantilla.exists():
            raise ValueError(f"Plantilla no encontrada: {plantilla}")
    else:
        plantilla = encontrar_plantilla(doc_meta)
    nombre_base = f"{codigo_final}_{_slug(doc_meta['nombre'])}"
    es_word = doc_meta["formato"] == "Word"
    nombre_archivo = f"{nombre_base}.docx" if es_word else f"{nombre_base}.xlsx"
    mime = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if es_word
        else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    with tempfile.TemporaryDirectory(prefix="sst_gen_") as tmp:
        out_path = Path(tmp) / nombre_archivo
        if es_word:
            if plantilla:
                generar_word_desde_plantilla(
                    empresa, doc_gen, version, fecha, plantilla, out_path
                )
                modo = "plantilla"
            else:
                generar_word_desde_cero(empresa, doc_gen, version, fecha, out_path)
                modo = "desde_cero"
        else:
            if plantilla:
                generar_excel_desde_plantilla(
                    empresa, doc_gen, version, fecha, plantilla, out_path
                )
                modo = "plantilla"
            else:
                generar_excel_desde_cero(empresa, doc_gen, version, fecha, out_path)
                modo = "desde_cero"
        # Pase final obligatorio: logo + empresa en Word y Excel
        try:
            if es_word:
                inyectar_encabezado_empresa_word(
                    out_path, empresa, doc_gen, version, fecha
                )
            else:
                asegurar_marca_excel(
                    out_path, empresa, doc_gen, version, fecha
                )
        except Exception:
            pass
        contenido = out_path.read_bytes()

    db.registrar_historial(codigo_final, doc_meta["nombre"], version, nombre_archivo, fecha)
    return {
        "codigo": codigo_final,
        "codigo_catalogo": codigo,
        "nombre": doc_meta["nombre"],
        "version": version,
        "fecha": fecha,
        "numero": n,
        "archivo": nombre_archivo,
        "nombre_archivo": nombre_archivo,
        "contenido": contenido,
        "mime": mime,
        "modo": modo,
        "plantilla": plantilla.name if plantilla else None,
    }
