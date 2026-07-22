# -*- coding: utf-8 -*-
"""Inyecta logo + datos de empresa en encabezados Word (todas las plantillas)."""
from __future__ import annotations

from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generador.colores_marca import colores_desde_logo


def _logo_path(empresa: dict) -> Optional[Path]:
    raw = (empresa.get("logo_path") or "").strip()
    if not raw:
        return None
    p = Path(raw)
    if not p.exists() or p.name == "logo.png":
        return None
    return p


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _clear_cell(cell) -> None:
    """Quita texto e imagenes de la celda, deja un parrafo vacio."""
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:tcPr"):
            continue
        tc.remove(child)
    tc.append(OxmlElement("w:p"))


def _force_black(run) -> None:
    try:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except Exception:
        pass


def _set_run_color(run, hex_color: str) -> None:
    try:
        run.font.color.rgb = RGBColor.from_string(hex_color)
    except Exception:
        _force_black(run)


def _meta_empresa_texto(
    empresa: dict,
    codigo: str,
    version: int | str,
    fecha: str,
) -> str:
    razon = (empresa.get("razon_social") or "").strip() or "Empresa no configurada"
    nit = (empresa.get("nit") or "").strip() or "____"
    lines = [
        razon,
        f"NIT: {nit}",
        f"Codigo: {codigo}",
        f"Version: {version}",
        f"Fecha: {fecha}",
    ]
    arl = (empresa.get("arl") or "").strip()
    if arl:
        lines.insert(2, f"ARL: {arl}")
    return "\n".join(lines)


def _poner_logo_en_celda(cell, logo: Optional[Path], fallback: str, fill_hex: str) -> None:
    _clear_cell(cell)
    _shade_cell(cell, fill_hex)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo:
        try:
            p.add_run().add_picture(str(logo), width=Inches(1.0))
            return
        except Exception:
            pass
    if fallback:
        run = p.add_run(fallback)
        run.bold = True
        run.font.size = Pt(8)
        _force_black(run)


def _poner_meta_en_celda(cell, texto: str, fill_hex: str, on_hex: str) -> None:
    _clear_cell(cell)
    _shade_cell(cell, fill_hex)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(texto)
    run.font.size = Pt(8)
    _set_run_color(run, on_hex)


def _header_tiene_empresa(header, razon: str) -> bool:
    if not razon:
        return False
    needle = razon.lower()
    for table in header.tables:
        for row in table.rows:
            for cell in row.cells:
                if needle in (cell.text or "").lower():
                    return True
    for p in header.paragraphs:
        if needle in (p.text or "").lower():
            return True
    return False


def _crear_tabla_encabezado(
    header,
    empresa: dict,
    titulo: str,
    codigo: str,
    version: int | str,
    fecha: str,
) -> None:
    logo = _logo_path(empresa)
    palette = colores_desde_logo(logo)
    # Limpiar parrafos vacios del header para no empujar la tabla
    for p in list(header.paragraphs):
        if not (p.text or "").strip():
            p._element.getparent().remove(p._element)

    table = header.add_table(rows=1, cols=3, width=Inches(7.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.autofit = True
    except Exception:
        pass

    razon = (empresa.get("razon_social") or "").strip()
    _poner_logo_en_celda(table.cell(0, 0), logo, razon[:40], palette["light"])

    title_cell = table.cell(0, 1)
    _clear_cell(title_cell)
    _shade_cell(title_cell, palette["primary"])
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pt = title_cell.paragraphs[0]
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = pt.add_run((titulo or "").upper())
    tr.bold = True
    tr.font.size = Pt(11)
    _set_run_color(tr, palette.get("on_primary", "243447"))

    _poner_meta_en_celda(
        table.cell(0, 2),
        _meta_empresa_texto(empresa, codigo, version, fecha),
        palette["secondary"],
        palette.get("on_primary", "243447"),
    )


def _actualizar_tabla_encabezado(
    table,
    empresa: dict,
    codigo: str,
    version: int | str,
    fecha: str,
) -> None:
    logo = _logo_path(empresa)
    palette = colores_desde_logo(logo)
    razon = (empresa.get("razon_social") or "").strip()
    ncols = len(table.columns)
    nrows = len(table.rows)
    if ncols < 1 or nrows < 1:
        return

    # Columna logo (primera): imagen o nombre
    _poner_logo_en_celda(table.cell(0, 0), logo, razon[:40], palette["light"])

    # Columna meta (ultima): empresa + nit + codigo/version/fecha
    meta_cell = table.cell(0, ncols - 1)
    _poner_meta_en_celda(
        meta_cell,
        _meta_empresa_texto(empresa, codigo, version, fecha),
        palette["secondary"],
        palette.get("on_primary", "243447"),
    )

    # Si hay mas filas en la columna logo (celdas fusionadas a veces se ven vacias),
    # no forzar texto en todas para no romper merges; la fila 0 basta.


def inyectar_encabezado_empresa_word(
    out_path: Path | str,
    empresa: dict,
    doc_meta: dict[str, Any],
    version: int | str,
    fecha: str,
) -> None:
    """Garantiza logo + razon social + NIT en el encabezado de cada seccion."""
    path = Path(out_path)
    if not path.exists():
        return
    doc = Document(str(path))
    codigo = str(doc_meta.get("codigo") or "")
    titulo = str(doc_meta.get("nombre") or "")
    razon = (empresa.get("razon_social") or "").strip()

    for section in doc.sections:
        header = section.header
        try:
            header.is_linked_to_previous = False
        except Exception:
            pass

        if header.tables:
            # Actualizar siempre la primera tabla (logo + meta),
            # aunque ya diga la razon social (para refrescar logo/codigo).
            _actualizar_tabla_encabezado(
                header.tables[0], empresa, codigo, version, fecha
            )
        else:
            _crear_tabla_encabezado(
                header, empresa, titulo, codigo, version, fecha
            )

        # Si tras actualizar aun no aparece la empresa (tabla rara), crear banner
        if razon and not _header_tiene_empresa(header, razon):
            p = header.add_paragraph()
            run = p.add_run(
                f"{razon}  |  NIT: {(empresa.get('nit') or '').strip() or '____'}  |  "
                f"{codigo}  |  v{version}  |  {fecha}"
            )
            run.bold = True
            run.font.size = Pt(9)
            _force_black(run)

    doc.save(str(path))
