# -*- coding: utf-8 -*-
"""Reparacion de layouts rotos en plantillas de encuesta."""
from __future__ import annotations

import re
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# Encuesta sociodemografica FG-FOR-SST-39: pares izquierda/derecha
_FILAS_SOCIO = [
    (
        ("1. EDAD", ["a. 18 – 25 años", "b. 26 – 35 años", "c. 36 – 45 años", "d. 46 años o más"]),
        (
            "2. ESTADO CIVIL",
            ["a. Soltero (a)", "b. Unión libre", "c. Casado (a)", "d. Divorciado (a)"],
        ),
    ),
    (
        ("3. SEXO", ["a. Hombre", "b. Mujer", "c. No binario"]),
        (
            "4. NÚMERO DE PERSONAS A CARGO",
            ["a. Ninguna", "b. 1 – 3 personas", "c. 4 – 6 personas", "d. Más de 6 personas"],
        ),
    ),
    (
        (
            "5. NIVEL DE ESCOLARIDAD",
            [
                "a. Primaria",
                "b. Secundaria",
                "c. Técnico / Tecnólogo",
                "d. Universitario",
                "e. Ninguna",
            ],
        ),
        (
            "6. TENENCIA DE VIVIENDA",
            [
                "a. Propia",
                "b. Arrendada",
                "c. Familiar",
                "d. Compartida con otra(s) familia(s)",
            ],
        ),
    ),
    (
        (
            "7. USO DEL TIEMPO LIBRE",
            [
                "a. Otro trabajo",
                "b. Labores domésticas",
                "c. Recreación y deporte",
                "d. Estudio",
                "e. Ninguno",
            ],
        ),
        (
            "8. PROMEDIO DE INGRESOS (S.M.L.)",
            [
                "a. Mínimo Legal (S.M.L.)",
                "b. Entre 1 a 2 S.M.L.",
                "c. Entre 3 a 4 S.M.L.",
                "d. Más de 5 S.M.L.",
            ],
        ),
    ),
    (
        (
            "9. ANTIGÜEDAD EN LA EMPRESA",
            [
                "a. Menos de 1 mes",
                "b. De 1 a 6 meses",
                "c. De 7 a 11 meses",
                "d. De 1 a 2 años",
                "e. Más de 2 años",
            ],
        ),
        (
            "10. ANTIGÜEDAD EN EL CARGO ACTUAL",
            [
                "a. Menos de 1 mes",
                "b. De 1 a 6 meses",
                "c. De 7 a 11 meses",
                "d. De 1 a 2 años",
                "e. Más de 2 años",
            ],
        ),
    ),
    (
        (
            "11. HA PARTICIPADO EN ACTIVIDADES DE SALUD REALIZADAS POR LA EMPRESA",
            [
                "a. Vacunación",
                "b. Salud Oral",
                "c. Exámenes de laboratorio y otros",
                "d. Exámenes médicos anuales",
                "e. Ninguna",
            ],
        ),
        (
            "12. LE HAN DIAGNOSTICADO ALGUNA ENFERMEDAD",
            ["a. Sí", "b. No", "Cuál: _______________________________"],
        ),
    ),
    (
        (
            "13. CONSUME BEBIDAS ALCOHÓLICAS",
            ["a. Sí", "b. No", "Frecuencia: _________________________"],
        ),
        (
            "14. PRACTICA ALGÚN DEPORTE",
            [
                "___ Semanal",
                "___ Quincenal",
                "___ Mensual",
                "___ Ocasional",
                "Cuál: _______________________________",
            ],
        ),
    ),
    (
        ("15. FUMA", ["a. Sí", "b. No", "Promedio: ____________________________"]),
        ("16. CONSUME SUSTANCIAS TÓXICAS", ["a. Sí", "b. No"]),
    ),
    (
        ("17. CONSUME MEDICAMENTOS", ["a. Sí", "b. No"]),
        (
            "18. ESTRATO",
            ["a. 1", "b. 2", "c. 3", "d. 4", "e. 5", "f. 6"],
        ),
    ),
    (
        (
            "19. TIPO DE CONTRATO",
            [
                "a. Término indefinido",
                "b. Término fijo",
                "c. Obra o labor",
                "d. Prestación de servicios",
                "e. Otro: ________________",
            ],
        ),
        ("", []),  # celda derecha vacia para cerrar fila
    ),
]


def _set_cell_border(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcBorders"):
            tcPr.remove(child)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _fill_question_cell(cell, titulo: str, opciones: list[str]) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if titulo:
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    for op in opciones:
        p2 = cell.add_paragraph(op)
        p2.paragraph_format.space_after = Pt(1)
        p2.paragraph_format.space_before = Pt(0)
        for r in p2.runs:
            r.font.size = Pt(10)
    _set_cell_border(cell)
    try:
        cell.width = Cm(8.2)
    except Exception:
        pass


def _find_broken_socio_cell(doc: Document):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.upper()
                if "1. EDAD" in t and "MARQUE CON UNA X" in t and (
                    "PARTICIPADO" in t or "11." in t
                ):
                    # Ya reparado = tiene tabla anidada con varias filas
                    if cell.tables and len(cell.tables[0].rows) >= 6:
                        return None
                    return cell
    return None


def reparar_encuesta_sociodemografica(doc: Document) -> bool:
    """Convierte el cuestionario de espacios a tabla 2 columnas limpia."""
    target = _find_broken_socio_cell(doc)
    if target is None:
        return False

    old = target.text
    nombre = ""
    cargo = ""
    fecha_line = ""
    m = re.search(r"NOMBRE:\s*([^\n]*)", old, re.I)
    if m:
        nombre = m.group(1).strip(" _")
    m = re.search(r"CARGO:\s*([^\n]*)", old, re.I)
    if m:
        cargo = m.group(1).strip(" _")
    m = re.search(r"Fecha:\s*([^\n]*)", old, re.I)
    if m:
        fecha_line = m.group(1).strip()

    target.text = ""
    p = target.paragraphs[0]
    r = p.add_run(
        f"NOMBRE: {nombre or '________________________________________________'}"
    )
    r.font.size = Pt(10)
    p2 = target.add_paragraph(
        f"CARGO:  {cargo or '________________________________________________'}"
    )
    for rr in p2.runs:
        rr.font.size = Pt(10)
    p3 = target.add_paragraph(f"Fecha: {fecha_line or '____/____/________'}")
    for rr in p3.runs:
        rr.font.size = Pt(10)
    p4 = target.add_paragraph("Marque con una X:")
    if p4.runs:
        p4.runs[0].bold = True
        p4.runs[0].font.size = Pt(10)

    grid = target.add_table(rows=len(_FILAS_SOCIO), cols=2)
    for i, (izq, der) in enumerate(_FILAS_SOCIO):
        row = grid.rows[i]
        _fill_question_cell(row.cells[0], izq[0], izq[1])
        _fill_question_cell(row.cells[1], der[0], der[1])
    return True


def reparar_layouts_word(path, fecha: Optional[str] = None) -> bool:
    """Abre un docx generado, repara layouts conocidos y guarda."""
    doc = Document(str(path))
    changed = reparar_encuesta_sociodemografica(doc)
    if changed and fecha:
        # actualizar fecha del formulario si quedo la de ejemplo
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.startswith("Fecha:") and "02/02/2026" in p.text:
                            p.text = f"Fecha: {fecha}"
    if changed:
        doc.save(str(path))
    return changed
